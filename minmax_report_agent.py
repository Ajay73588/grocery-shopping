import os
import argparse
import re
import urllib.request
import tempfile
from urllib.parse import urlparse
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import wikipedia
from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup

def _clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text or '').strip()
    return text

def _extract_page_text(html: str) -> str:
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(['script', 'style', 'noscript', 'header', 'footer', 'nav', 'aside']):
        tag.decompose()
    article = soup.find('article')
    container = article if article else soup
    paragraphs = container.find_all('p')
    text_chunks = [_clean_text(p.get_text()) for p in paragraphs]
    text_chunks = [t for t in text_chunks if len(t) > 40]
    return "\n".join(text_chunks[:12])

def _domain_from_url(url: str) -> str:
    try:
        return urlparse(url).netloc.lower().lstrip('www.')
    except Exception:
        return ""

def do_research(topic: str) -> tuple[str, list[str]]:
    """Perform multi-source research using DDGS + Wikipedia fallback, emphasizing domain diversity."""
    print(f"[*] Deep Researching topic: {topic}")

    research_content = ""
    image_urls = []

    try:
        ddgs = DDGS()
        candidate_urls = []

        # 1) Text sources from multiple queries to avoid dependency on one site
        queries = [
            topic,
            f"{topic} research",
            f"{topic} analysis",
            f"{topic} report",
        ]
        for q in queries:
            try:
                results = list(ddgs.text(q, max_results=6))
                for res in results:
                    url = res.get('href', '')
                    title = res.get('title', '')
                    if url:
                        candidate_urls.append((url, title))
            except Exception:
                continue

        # 2) Add news sources for topical variety
        try:
            news_results = list(ddgs.news(topic, max_results=6))
            for res in news_results:
                url = res.get('url', '')
                title = res.get('title', '')
                if url:
                    candidate_urls.append((url, title))
        except Exception:
            pass

        # Dedupe by URL and prioritize domain diversity
        seen_urls = set()
        used_domains = set()
        sources_collected = 0
        max_sources = 7

        for url, title in candidate_urls:
            if url in seen_urls:
                continue
            seen_urls.add(url)
            domain = _domain_from_url(url)
            if domain and domain in used_domains and sources_collected < max_sources:
                continue

            try:
                req = requests.get(url, timeout=8, headers={'User-Agent': 'Mozilla/5.0'})
                if req.status_code != 200:
                    continue
                content_type = req.headers.get('Content-Type', '')
                if 'text/html' not in content_type:
                    continue

                text_snippet = _extract_page_text(req.text)
                if len(text_snippet) < 200:
                    continue

                title = title or _clean_text(BeautifulSoup(req.text, 'html.parser').title.string if BeautifulSoup(req.text, 'html.parser').title else "")
                research_content += f"Source: {title or domain}\nURL: {url}\nContent:\n{text_snippet}\n\n"
                used_domains.add(domain)
                sources_collected += 1
                if sources_collected >= max_sources:
                    break
            except Exception as scrape_e:
                print(f"[*] Failed to scrape URL {url}: {scrape_e}")
                continue

        # 3) Fetch images (filter to JPG/PNG for docx compatibility)
        try:
            img_results = list(ddgs.images(topic, max_results=10))
            for img in img_results:
                url = img.get('image', '')
                if not url:
                    continue
                if not url.lower().endswith(('.jpg', '.jpeg', '.png')):
                    continue
                if url not in image_urls:
                    image_urls.append(url)
                if len(image_urls) >= 5:
                    break
        except Exception:
            pass

        if not research_content.strip():
            raise Exception("No textual content successfully extracted.")

        return research_content, image_urls

    except Exception as e:
        print(f"[*] Primary research via DDGS failed ({str(e)[:50]}). Falling back to Wikipedia...")

        # Fallback Protocol: Multiple Wikipedia Pages
        try:
            search_results = wikipedia.search(topic, results=3)
            if not search_results:
                return "No significant research findings could be retrieved.", []

            for res in search_results:
                try:
                    page = wikipedia.page(res, auto_suggest=False)
                    research_content += f"Source: Wikipedia - {page.title}\nURL: {page.url}\nSummary:\n{page.summary}\n\n"
                    # Try to snag images from the pages if we still need them
                    for img in page.images:
                        if img.lower().endswith(('.png', '.jpg', '.jpeg')) and 'icon' not in img.lower() and 'logo' not in img.lower():
                            if 'commons-logo' not in img.lower() and 'wikiquote' not in img.lower():
                                if img not in image_urls:
                                    image_urls.append(img)
                        if len(image_urls) >= 3:
                            break
                except wikipedia.exceptions.DisambiguationError:
                    continue
                except Exception:
                    continue

            return research_content, image_urls[:3]

        except Exception as wiki_e:
            print(f"[*] Wikipedia fallback also failed: {wiki_e}")
            return "Research absolutely failed. Rely on internal reasoning.", []

def generate_report_content(topic: str, research_data: str, api_key: str, model_name: str, image_urls: list[str]) -> str:
    """Use MinMax via OpenAI-compatible API to generate a structured report."""
    print(f"[*] Generating report content with MinMax LLM (model: {model_name})...")
    
    client = OpenAI(
        api_key=api_key,
        base_url="https://api.minimax.io/v1", 
    )
    
    image_context = ""
    if image_urls:
        image_context = "\n=== AVAILABLE IMAGES ===\nYou have the following images available to embed into your report based on contextual relevance. You MUST use ALL image tags exactly once, each on its own line, and place them in contextually relevant sections.\n"
        for idx, url in enumerate(image_urls):
            filename = url.split('/')[-1]
            image_context += f"- Tag: [IMAGE_{idx+1}] - Filename Context: {filename}\n"
        image_context += "========================\n"

    prompt = f"""
You are an elite, professional research analyst. Your task is to craft a comprehensive, well-structured, and authoritative report on the topic: '{topic}'.

Base your factual claims and analysis on the following newly acquired research data:

=== RESEARCH DATA ===
{research_data}
=====================
{image_context}
Your report MUST strictly adhere to the following clean structure:

# {topic}

## Executive Summary
Provide a concise, high-level overview of the most critical takeaways. Keep it impactful.

## Introduction
Set the context, define the subject matter, and explain the significance of the topic.

## Comprehensive Analysis
Dive deep into the factual findings. 
- Break this section down using clear subheadings (e.g., `###`). 
- Present the information logically, using bullet points for lists if necessary.

## Strategic Conclusion
Summarize the analysis and offer a final perspective, summary, or future outlook.

## References
List all sources and URLs provided in the research data using a neat, bulleted format.

CRITICAL FORMATTING RULES:
1. Write in a formal, highly professional, and academic yet accessible tone.
2. Use markdown headings (`#`, `##`, `###`) for all sections to ensure proper document formatting.
3. Use bold text (`**text**`) to emphasize key terms.
4. DO NOT wrap your response in markdown code blocks (e.g., no ``` at the start or end).
5. DO NOT use markdown horizontal rules (like `---` or `***`).
6. Place the image tags (e.g., [IMAGE_1], [IMAGE_2]) strategically throughout the text on their own lines.
7. Output raw, clean text ready for document publication.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "You are a professional report generation AI."},
            {"role": "user", "content": prompt}
        ]
    )
    
    content = response.choices[0].message.content
    # Reason models sometimes output <think> blocks; remove them so they don't appear in the document
    content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()

    # Ensure all image tags are present so the docx embeds every image
    if image_urls:
        missing_tags = []
        for i in range(len(image_urls)):
            tag = f"[IMAGE_{i+1}]"
            if tag not in content:
                missing_tags.append(tag)
        if missing_tags:
            content += "\n\n## Image Appendix\n"
            for tag in missing_tags:
                content += f"{tag}\n"

    return content

def _add_markdown_paragraph(doc: Document, line: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = line.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = p.add_run(part)
            run.bold = True
        else:
            p.add_run(part)

def save_to_docx(topic: str, content: str, output_path: str, image_urls: list[str] = None):
    """Save the text content to a structured DOCX file."""
    print(f"[*] Saving report to {output_path}...")
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 1. Add Main Centered Title
    title = doc.add_heading(topic.upper(), 0)
    title.alignment = 1  # 1 = Center alignment
            
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        # Skip empty lines and markdown horizontal rules
        if not line or line.strip('-*_') == '':
            continue
            
        # Check if the line is a markdown heading (# to ######)
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            
            # If the LLM generates a H1 that directly matches the topic, skip it
            if level == 1 and text.strip().lower() == topic.lower():
                continue
                
            doc.add_heading(text, level)
            continue
            
        # Check for image tags anywhere in the line
        if '[IMAGE_' in line and ']' in line:
            if image_urls:
                try:
                    matches = list(re.finditer(r'\[IMAGE_(\d+)\]', line))
                    last_idx = 0
                    for match in matches:
                        # Add text before tag, if any
                        pre_text = line[last_idx:match.start()].strip()
                        if pre_text:
                            _add_markdown_paragraph(doc, pre_text)

                        idx = int(match.group(1)) - 1
                        if 0 <= idx < len(image_urls):
                            url = image_urls[idx]
                            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                            with urllib.request.urlopen(req) as response:
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                                    tmp_file.write(response.read())
                                    tmp_file_path = tmp_file.name

                            doc.add_picture(tmp_file_path, width=Inches(4.5))
                            doc.paragraphs[-1].alignment = 1
                            os.unlink(tmp_file_path)

                        last_idx = match.end()

                    # Add trailing text after last tag, if any
                    post_text = line[last_idx:].strip()
                    if post_text:
                        _add_markdown_paragraph(doc, post_text)
                except Exception as e:
                    print(f"[*] Failed to embed image placeholder {line}: {e}")
            continue

        # Handle basic bold markdown conversion for docx (simplistic)
        _add_markdown_paragraph(doc, line)
                    
    doc.save(output_path)
    print(f"[+] Successfully saved report as {output_path}")

def main():
    load_dotenv()  # Load environment variables from .env file
    parser = argparse.ArgumentParser(description="MinMax AI Report Generator")
    parser.add_argument("--topic", type=str, required=True, help="The topic for the report.")
    parser.add_argument("--output", type=str, default="Report.docx", help="Output docx file name.")
    parser.add_argument("--model", type=str, default="MiniMax-M2.7", help="The MinMax model ID to use.")
    args = parser.parse_args()

    api_key = os.environ.get("MINIMAX_API_KEY")
    if not api_key:
        print("Error: MINIMAX_API_KEY environment variable not set.")
        print("Please set it before running the script: set MINIMAX_API_KEY=your_key_here")
        return

    research_data, image_urls = do_research(args.topic)
    report_content = generate_report_content(args.topic, research_data, api_key, args.model, image_urls)
    save_to_docx(args.topic, report_content, args.output, image_urls)

if __name__ == "__main__":
    main()
